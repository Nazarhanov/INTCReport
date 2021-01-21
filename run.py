import yaml
import docx
import sys
import re


def main():
    Report = docx.Document("templates/ru.edu.docx")

    ReportTitle = Report.sections[0]
    ReportContent = Report.sections[1]
    ReportBody = Report.sections[2]

    # get args
    configPath = sys.argv[-1]
    outPath = '.'.join(configPath.split(".")[:-1]) + '.out.docx'

    # load yaml
    with open(configPath, encoding="utf-8") as config:
        try:
            config = yaml.load(config, Loader=yaml.Loader)
        except:
            raise Exception("Error on reading yaml file")

    # replace @ variables in Body Paragraphs
    title = {
        "@subject": config["title"]["subject"],
        "@author": f'{config["title"]["author"]["surname"]} {config["title"]["author"]["name"]} {config["title"]["author"]["patronymic"]}',
        "@course": config["title"]["course"],
        "@group": config["title"]["group"],
        "@speciality.id": config["title"]["speciality"]["id"],
        "@speciality.name": config["title"]["speciality"]["name"],
        "@qualification.id": config["title"]["qualification"]["id"],
        "@qualification.name": config["title"]["qualification"]["name"],
        "@place": config["title"]["place"],
        "@period.start.day": config["title"]["period"]["start"]["day"],
        "@period.start.month": config["title"]["period"]["start"]["month"],
        "@period.start.year": config["title"]["period"]["start"]["year"],
        "@period.end.day": config["title"]["period"]["end"]["day"],
        "@period.end.month": config["title"]["period"]["end"]["month"],
        "@period.end.year": config["title"]["period"]["end"]["year"],
        "@mentor": f'{config["title"]["mentor"]["surname"]} {config["title"]["mentor"]["name"]} {config["title"]["mentor"]["patronymic"]}',
        "@mentor.suffix": config["title"]["mentor"]["suffix"],
    }

    def replaceVars(variables, paragraphs):
        keys = list(variables)

        for paragraph in paragraphs:
            if "@" in paragraph.text:
                paragraph = paragraph.runs
                for index in range(len(paragraph)):
                    word = paragraph[index].text
                    if vs := [ v.group() for v in re.finditer('@[A-Za-z.]+', word)]:
                        for word in vs:
                            if word in keys:
                                newParagraph = paragraph[index].text.replace(word, str(variables[word]))
                                paragraph[index].text = newParagraph
                    # if ("@" in word) and (word in keys):
                    #     paragraph[index].text = str(variables[word])

    # replace @ variables in Content and Body Footer frames
    frame = {
        "@subject": config["title"]["subject"],
        "@author": f'{config["title"]["author"]["surname"]} {config["title"]["author"]["name"][0]}.{config["title"]["author"]["patronymic"][0]}.',
        "@mentor": f'{config["title"]["mentor"]["surname"]} {config["title"]["mentor"]["name"][0]}.{config["title"]["mentor"]["patronymic"][0]}.',
        "@group": config["title"]["group"],
        "@speciality.id": config["title"]["speciality"]["id"],
    }

    for row in ReportContent.first_page_footer.tables[0].rows:
        for cell in row.cells:
            replaceVars(variables=frame, paragraphs=cell.paragraphs)

    for row in ReportBody.footer.tables[0].rows:
        for cell in row.cells:
            replaceVars(variables=frame, paragraphs=cell.paragraphs)

    #
    def add_paragraph_with_bookmark(paragraph, text, id):
        id = str(id)
        te = str(text)

        run = paragraph.add_run()
        tag = run._r
        start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
        start.set(docx.oxml.ns.qn('w:id'), '0')
        start.set(docx.oxml.ns.qn('w:name'), id)
        tag.append(start)

        text = docx.oxml.OxmlElement('w:r')
        text.text = te
        tag.append(text)

        end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
        end.set(docx.oxml.ns.qn('w:id'), '0')
        end.set(docx.oxml.ns.qn('w:name'), id)
        tag.append(end)

        return paragraph

    # See more here
    # https://github.com/python-openxml/python-docx/pull/210
    # https://gist.github.com/photuris/8f71166deec7f8f5a774
    def reset_list(paragraph, num_id):
        p_xml = paragraph._p
        p_props = p_xml.get_or_add_pPr()
        num_props = docx.oxml.shared.OxmlElement('w:numPr')

        lvl_prop = docx.oxml.shared.OxmlElement('w:ilvl')
        lvl_prop.set(docx.oxml.ns.qn('w:val'), '0')

        num_id_prop = docx.oxml.shared.OxmlElement('w:numId')
        num_id_prop.set(docx.oxml.ns.qn('w:val'), str(num_id))

        num_props.append(lvl_prop)
        num_props.append(num_id_prop)

        p_props.append(num_props)

    def get_renumbered_num_id(document, num_id):
        numbering = document._part \
                            .numbering_part \
                            .numbering_definitions \
                            ._numbering

        next_num_id = numbering._next_numId

        num = docx.oxml.numbering.CT_Num.new(next_num_id, num_id)
        num.add_lvlOverride(ilvl=0).add_startOverride(1)

        new_num = numbering._insert_num(num)

        return new_num.numId
    #
    firstH1 = True
    firstH2 = True

    idH1 = 0
    idIMG = 1
    idBOOKMARKS = 0
    noBOOKMARDS = [] # only for first table column with h1c

    idNumList = get_renumbered_num_id(Report, 10)
    firstNext = True

    # paste Body of report
    for paragraph in config["body"]:
        key = list(paragraph)[0]

        # before an element will paste
        if key == "h1c":
            noBOOKMARDS.append(idBOOKMARKS)

        if key == "h1c" or key == "h1":
            if firstH1:
                firstH1 = False
            else:
                Report.add_page_break()

        elif key == "h2":
            if firstH2:
                firstH2 = False
            else:
                Report.add_page_break()

        elif key == "h3":
            Report.add_paragraph().style = Report.styles["space"]

        # while an element is being inserted
        if key == "img":
            val = paragraph[key]
            size = val[2] if 2 < len(val) else None

            src = val[0]

            if size == "small":
                Report.add_picture(src, width=docx.shared.Inches(2))
            else: 
                Report.add_picture(src, width=docx.shared.Inches(4.5))
            
            Report.paragraphs[-1].style = Report.styles["img"]
            
            text = f"Рис. {idH1}.{idIMG} {val[1]}"
            Report.add_paragraph(text).style = Report.styles["subtitle"]

            idIMG += 1
        elif key == "h1c" or key == "h1" or key == "h2" or key == "h3":
            newParagraph = Report.add_paragraph()
            newParagraph = add_paragraph_with_bookmark(
                paragraph=newParagraph, text=paragraph[key], id=idBOOKMARKS)
            newParagraph.style = Report.styles[key]

            idBOOKMARKS += 1
        elif key == "items":
            for item in paragraph[key]:
                newParagraph = Report.add_paragraph(item)
                newParagraph.style = Report.styles["item"]
        elif key == "steps":
            if type(paragraph[key]) == list:
                idNumList = get_renumbered_num_id(Report, 10)
                for item in paragraph[key]:
                    newStep = Report.add_paragraph(item)
                    newStep.style = Report.styles["step"]
                    reset_list(paragraph=newStep, num_id=idNumList)
            elif type(paragraph[key]) == str:
                if firstNext:
                    idNumList = get_renumbered_num_id(Report, 10)
                    firstNext = False
                newStep = Report.add_paragraph(paragraph[key])
                newStep.style = Report.styles["step"]
                reset_list(paragraph=newStep, num_id=idNumList)

        elif key == "steps-end":
            newStep = Report.add_paragraph(paragraph[key])
            newStep.style = Report.styles["step"]
            reset_list(paragraph=newStep, num_id=idNumList)
            firstNext = True
        else:
            newParagraph = Report.add_paragraph(text=paragraph[key])
            newParagraph.style = Report.styles[key]

        # after an element will paste
        if key == "h1c" or key == "h2" or key == "h3":
            Report.add_paragraph().style = Report.styles["space"]
            Report.add_paragraph().style = Report.styles["space"]
        elif key == "h1":
            Report.add_paragraph().style = Report.styles["space"]
            firstH2 = True
            idH1 += 1
            idIMG = 1

    #
    def add_ref_to_bookmark(paragraph, refId, refType=0):
        runP = paragraph.add_run()
        r = runP._r
        fldChar = docx.oxml.shared.OxmlElement('w:fldChar')
        fldChar.set(docx.oxml.ns.qn('w:fldCharType'), 'begin')
        r.append(fldChar)
        instrText = docx.oxml.shared.OxmlElement('w:instrText')

        if refType == 1:
            refType = "REF"
            refSpec = "\w \h"
        elif refType == 2:
            refType = "PAGEREF"
            refSpec = "\h"
        else:
            refType = "REF"
            refSpec = "\n"
        instrText.text = f" {refType} {refId} {refSpec} "
        r.append(instrText)
        fldChar = docx.oxml.shared.OxmlElement('w:fldChar')
        fldChar.set(docx.oxml.ns.qn('w:fldCharType'), 'end')
        r.append(fldChar)

    #
    firstRow = True
    ContentTable = Report.tables[0]

    # create Content
    for i in range(idBOOKMARKS):

        if firstRow:
            newCells = ContentTable.rows[0].cells
            firstRow = False
        else:
            newCells = ContentTable.add_row().cells

        if not(i in noBOOKMARDS):
            add_ref_to_bookmark(
                paragraph=newCells[0].paragraphs[0], refId=i, refType=1)
        add_ref_to_bookmark(
            paragraph=newCells[1].paragraphs[0], refId=i, refType=0)
        add_ref_to_bookmark(
            paragraph=newCells[2].paragraphs[0], refId=i, refType=2)

        newCells[2].paragraphs[0].style = Report.styles["page number"]

    # start replacing variables in body paragraphs
    replaceVars(variables=title, paragraphs=Report.paragraphs)

    Report.save(outPath)


if __name__ == "__main__":
    main()
